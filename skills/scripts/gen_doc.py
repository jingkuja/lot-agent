#!/usr/bin/env python3
"""Generate a document file from a JSON spec.

Usage:
    python gen_doc.py <spec.json>

Spec (JSON object):
    {
      "title":   "Document title",          # optional
      "content": "Markdown body text",       # required
      "format":  "docx" | "pdf" | "md" | "html",
      "outPath": "/abs/path/to/output.docx"  # extension matches `format`
    }

Prints a single JSON line to stdout:
    {"ok": true, "path": "...", "format": "docx", "degraded": false, "note": ""}

If the third-party library for the requested format is unavailable, the script
DEGRADES to Markdown (writes a .md next to outPath), and reports degraded=true.
Only the Python standard library is required for md/html; docx needs
python-docx and pdf needs reportlab.
"""
import json
import os
import re
import sys
import html as html_mod


# ── tiny markdown parser (headings / bullets / paragraphs) ──────────────────

def parse_blocks(content):
    """Yield (kind, text) blocks. kind in {h1,h2,h3,bullet,para,blank}."""
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line.strip():
            yield ("blank", "")
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            yield ("h" + str(len(m.group(1))), m.group(2).strip())
            continue
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if m:
            yield ("bullet", m.group(1).strip())
            continue
        yield ("para", line.strip())


def inline_plain(text):
    """Strip basic inline markdown (**bold**, *italic*, `code`)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


# ── writers ─────────────────────────────────────────────────────────────────

def write_md(title, content, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        if title:
            f.write("# %s\n\n" % title)
        f.write(content)
        if not content.endswith("\n"):
            f.write("\n")
    return out_path


def write_html(title, content, out_path):
    parts = [
        "<!DOCTYPE html>",
        '<html lang="zh-CN"><head><meta charset="utf-8">',
        "<title>%s</title>" % html_mod.escape(title or "Document"),
        "<style>body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;"
        "max-width:760px;margin:40px auto;padding:0 20px;line-height:1.7;"
        "color:#1a1a22}h1,h2,h3{line-height:1.3}code{background:#f3f3f7;"
        "padding:2px 5px;border-radius:4px}</style></head><body>",
    ]
    if title:
        parts.append("<h1>%s</h1>" % html_mod.escape(title))
    in_list = False
    for kind, text in parse_blocks(content):
        esc = html_mod.escape(inline_plain(text))
        if kind == "bullet":
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append("<li>%s</li>" % esc)
            continue
        if in_list:
            parts.append("</ul>")
            in_list = False
        if kind in ("h1", "h2", "h3"):
            parts.append("<%s>%s</%s>" % (kind, esc, kind))
        elif kind == "para":
            parts.append("<p>%s</p>" % esc)
    if in_list:
        parts.append("</ul>")
    parts.append("</body></html>")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))
    return out_path


def write_docx(title, content, out_path):
    from docx import Document  # python-docx

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for kind, text in parse_blocks(content):
        plain = inline_plain(text)
        if kind == "h1":
            doc.add_heading(plain, level=1)
        elif kind == "h2":
            doc.add_heading(plain, level=2)
        elif kind == "h3":
            doc.add_heading(plain, level=3)
        elif kind == "bullet":
            doc.add_paragraph(plain, style="List Bullet")
        elif kind == "para":
            doc.add_paragraph(plain)
    doc.save(out_path)
    return out_path


def write_pdf(title, content, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        ListFlowable,
        ListItem,
    )

    styles = getSampleStyleSheet()
    story = []
    if title:
        story.append(Paragraph(html_mod.escape(title), styles["Title"]))
        story.append(Spacer(1, 12))
    bullets = []

    def flush_bullets():
        if bullets:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(b, styles["BodyText"])) for b in bullets],
                    bulletType="bullet",
                )
            )
            bullets.clear()

    for kind, text in parse_blocks(content):
        esc = html_mod.escape(inline_plain(text))
        if kind == "bullet":
            bullets.append(esc)
            continue
        flush_bullets()
        if kind in ("h1", "h2", "h3"):
            level = {"h1": "Heading1", "h2": "Heading2", "h3": "Heading3"}[kind]
            story.append(Paragraph(esc, styles[level]))
        elif kind == "para":
            story.append(Paragraph(esc, styles["BodyText"]))
        elif kind == "blank":
            story.append(Spacer(1, 6))
    flush_bullets()
    SimpleDocTemplate(out_path, pagesize=A4).build(story)
    return out_path


WRITERS = {"md": write_md, "html": write_html, "docx": write_docx, "pdf": write_pdf}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "note": "missing spec path"}))
        return 1
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        spec = json.load(f)

    title = spec.get("title", "") or ""
    content = spec.get("content", "") or ""
    fmt = spec.get("format", "docx")
    out_path = spec["outPath"]

    if fmt not in WRITERS:
        fmt = "docx"

    degraded = False
    note = ""
    try:
        WRITERS[fmt](title, content, out_path)
        produced = out_path
    except ImportError as exc:
        # Library for the requested format is not installed → degrade to Markdown.
        degraded = True
        note = "missing library for %s (%s); produced Markdown instead" % (fmt, exc)
        produced = os.path.splitext(out_path)[0] + ".md"
        write_md(title, content, produced)
        fmt = "md"

    print(
        json.dumps(
            {"ok": True, "path": produced, "format": fmt, "degraded": degraded, "note": note}
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
